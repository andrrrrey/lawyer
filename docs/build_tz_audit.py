from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("Отчёт-по-соответствию-ТЗ-Lawyer.docx")
BLUE = "315E86"
DARK = "1F2937"
MUTED = "667085"
LINE = "D0D5DD"
LIGHT = "F2F4F7"
GREEN = "E7F6EC"
YELLOW = "FFF4CC"
RED = "FDECEC"
PURPLE = "F1ECFF"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_fixed(table) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_border(cell, color=LINE, size="4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_run(run, *, size=None, bold=None, color=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char1, instr, fld_char2))
    set_run(run, size=8.5, color=MUTED)


def status_fill(status: str) -> str:
    if status == "Реализовано":
        return GREEN
    if status in {"Частично", "Требует настройки"}:
        return YELLOW
    if status == "Отложено":
        return PURPLE
    return RED


def add_para(doc, text: str = "", *, style=None, bold_lead: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def add_bullet(doc, text: str, level: int = 0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    return add_para(doc, text, style=style)


def add_number(doc, text: str):
    return add_para(doc, text, style="List Number")


def add_callout(doc, title: str, text: str, fill=LIGHT, trailing=True) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    cell = table.cell(0, 0)
    set_cell_width(cell, 9360)
    set_cell_margins(cell, 130, 180, 130, 180)
    set_cell_border(cell, color=fill, size="1")
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run(r, bold=True, color=BLUE, size=10.5)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run(r, size=9.5, color=DARK)
    if trailing:
        doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_matrix(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    widths = (1950, 1320, 4650, 1440)
    headers = ("Требование ТЗ", "Статус", "Как реализовано / где находится", "Что осталось")
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, cell in enumerate(hdr.cells):
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell)
        set_cell_border(cell)
        shade(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(headers[idx])
        set_run(r, size=8.3, bold=True, color=DARK)
    for req, status, impl, remain in rows:
        row = table.add_row()
        set_row_cant_split(row)
        cells = row.cells
        values = (req, status, impl, remain)
        for idx, cell in enumerate(cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if idx == 1:
                shade(cell, status_fill(status))
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(values[idx])
            set_run(r, size=8.2, bold=(idx in (0, 1)), color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_two_col(doc, rows, widths=(2600, 6760)):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    for label, value in rows:
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(label if idx == 0 else value)
            set_run(r, size=9, bold=(idx == 0), color=BLUE if idx == 0 else DARK)
        shade(cells[0], LIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(4)
    for name, size, before, after in (("Title", 24, 0, 12), ("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6), ("Heading 3", 12, 8, 4)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BLUE if name != "Title" else DARK)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("LAWYER  /  АУДИТ ТЕХНИЧЕСКОГО ЗАДАНИЯ")
    set_run(r, size=8.5, bold=True, color=MUTED)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("Конфиденциально  ·  05.09.2026  ·  ")
    set_run(r, size=8.5, color=MUTED)
    add_page_field(fp)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("ОТЧЁТ О СООТВЕТСТВИИ ТЗ")
    set_run(r, size=10, bold=True, color=BLUE)
    p = doc.add_paragraph(style="Title")
    r = p.add_run("Lawyer")
    set_run(r, size=24, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Управленческая панель продаж и маркетинга")
    set_run(r, size=16, color=MUTED)
    add_two_col(doc, [
        ("Основание", "«ТЗ — панель продаж — итоговое — Lawyer»"),
        ("Объект проверки", "Исходный код, база данных и production lawyer.futuguru.com"),
        ("Дата среза", "5 сентября 2026 года"),
        ("Назначение", "Сверка требований, текущей реализации и остаточных работ"),
    ])
    add_callout(doc, "Главный вывод", "Функциональные доработки по семи пунктам выполнены: загружаются история стадий и активности, добавлены глобальный фильтр воронки, план-факт, ожидаемые оплаты, средние суммы, цикл сделки, аналитика отделов, роли и комментарии по сделкам. До приёмки остаётся бизнес-настройка: выбрать 11 воронок, выдать Bitrix24 право user_brief, заполнить сотрудников, отделы и планы, а также завершить параметры трёх аккаунтов Яндекса и проверить факты 1С.", fill=YELLOW)

    doc.add_heading("1. Как читать отчёт", level=1)
    add_two_col(doc, [
        ("Реализовано", "Функция есть в коде, видна в интерфейсе или работает в сервисном слое."),
        ("Частично", "Есть каркас или часть расчёта, но не все поля, формулы или источники доведены до результата."),
        ("Требует настройки", "Код готов, но нужны доступы, сопоставления или начальные данные Заказчика."),
        ("Не реализовано", "Требуется доработка кода."),
        ("Отложено", "Исключено из текущего запуска по прямому указанию Заказчика."),
    ])

    doc.add_heading("2. Итог аудита", level=1)
    add_para(doc, "Проверка выполнена по актуальному ТЗ, коду, миграциям, API, production-базе и собранному web-интерфейсу. Доступы и секреты в отчёт не включены.")
    add_bullet(doc, "Платформа развёрнута на https://lawyer.futuguru.com, API в production отвечает, PostgreSQL и миграции актуальны.")
    add_bullet(doc, "В production 779 сделок Bitrix24, 2 360 переходов стадий и 1 999 CRM-активностей. История, звонки и встречи поступают из обоих порталов.")
    add_bullet(doc, "Комментарии сформированы для 779 сделок. Базовый комментарий строится из фактов CRM; LLM-версия включается после настройки AI-интеграции.")
    add_bullet(doc, "Поступлений 1С, расходов Директа, визитов Метрики и ручных расходов пока нет. Финансовые и ROMI-блоки корректно остаются без расчётных значений до появления фактов.")
    add_bullet(doc, "Включено 15 воронок вместо 11 из ТЗ; это не ошибка механизма, но незавершённая бизнес-настройка.")
    add_bullet(doc, "Отделы, сотрудники, планы и отдельные пользователи панели в production не заполнены. Вход выполняется резервной учётной записью собственника.")
    add_bullet(doc, "Оба webhook Bitrix24 читают CRM, но отклоняют user.get из-за недостаточных прав; поэтому часть ответственных пока отображается по ID.")
    add_bullet(doc, "Исходный код опубликован в GitHub main, актуальный коммит 4f13e8d. Последнее исправление прошло lint и 42 целевых теста.")

    doc.add_heading("3. Источники данных и архитектура", level=1)
    add_matrix(doc, [
        ("Один дашборд; 3 юрлица", "Реализовано", "Единое приложение. Глобальный фильтр юрлица; Настройки → Юрлица. ИНН ЮО, ЦСВ и УрПАСЭ заданы.", "Ничего по трём основным юрлицам."),
        ("2 портала Bitrix24", "Реализовано", "Два входящих webhook. Сделки, история стадий, звонки и встречи загружаются раздельно по источникам box/cloud. В production: 779 сделок, 2 360 переходов и 1 999 активностей.", "Добавить право user_brief для загрузки ФИО."),
        ("Имена ответственных", "Требует настройки", "Автоматический справочник user.get и ручное сопоставление ID → ФИО реализованы. Ручное сохранение сразу обновляет ранее загруженные сделки.", "В обоих webhook включить scope user_brief либо заполнить Настройки → Сотрудники вручную."),
        ("1С: POST + Basic", "Требует настройки", "Адаптер боевого endpoint, HTTP POST, Basic Authentication, скользящее окно, дедупликация и журнал сопоставлений реализованы.", "В production 0 поступлений: проверить доступность endpoint и учётные данные, затем запустить пересчёт."),
        ("Фильтры поступлений 1С", "Реализовано", "ИНН получателя; Код_BTX/Тип_BTX; исключение внутригрупповых переводов; белый список точных статей ДДС; типы операций.", "Тест на боевой выгрузке и разбор исключённых/несопоставленных строк."),
        ("РЦА", "Отложено", "Не включено в три утверждённых юрлица.", "Нужен точный ИНН и решение о включении."),
        ("3 аккаунта Директа/Метрики", "Требует настройки", "Интеграции → Яндекс: отдельные комплекты настроек для ЮО, ЦСВ и УрПАСЭ; адаптеры Direct Reports API и Metrika API.", "OAuth-токены сохранены. Для каждого юрлица заполнить логин Директа и ID счётчика Метрики; в базе пока 0 расходов и 0 визитов."),
    ])

    doc.add_heading("4. Воронка, планирование и менеджеры", level=1)
    add_matrix(doc, [
        ("11 воронок", "Требует настройки", "Настройки → Воронки: система читает их из обоих Bitrix24, позволяет выбрать нужные, привязать юрлицо и SLA.", "Сейчас включено 15. Оставить точно 7 ЮО + 2 УрПАСЭ + 2 lead-воронки."),
        ("Этапы и конверсии", "Реализовано", "Воронка обработки использует фактическую историю стадий. Глобальный фильтр воронки применяется к KPI, таблицам и аналитическим блокам.", "Провести бизнес-приёмку после выбора окончательных 11 воронок."),
        ("Планы по месяцам", "Требует настройки", "План-факт реализован для компании, отдела и сотрудника: выручка, оплаты, сделки, звонки и встречи, факт и процент выполнения.", "В production планов 0. Заполнить структуру и месячные планы."),
        ("Ожидание оплат", "Реализовано", "Отдельный расчёт ожидаемых оплат по настроенным стадиям, суммам сделок, периоду, юрлицу и воронке.", "Подтвердить целевые стадии, включая «Заключение Контракта»."),
        ("Средние суммы", "Реализовано", "На дашборде рассчитаны средняя сумма договора по CRM и средняя сумма поступления по фактам 1С.", "Средняя сумма поступления появится после загрузки 1С."),
        ("Цикл сделки", "Реализовано", "Средний цикл рассчитывается по фактическим переходам и датам сделок с учётом глобальных фильтров.", "Проверить трактовку финальной стадии на выбранных воронках."),
        ("Карточки менеджеров", "Реализовано", "Контроль менеджеров дополнен конверсией и план-фактом продаж, звонков и встреч.", "Настроить ФИО и соответствие сотрудников, затем внести планы."),
        ("Агрегаты отделов/МОП", "Требует настройки", "Добавлен отдельный блок отделов: сотрудники, лиды, сделки в работе, конверсия, звонки, встречи, оплаты и выручка.", "В production отделов 0; создать отделы и привязать сотрудников."),
    ])

    doc.add_heading("5. Контроль, SLA и AI", level=1)
    add_matrix(doc, [
        ("Сделки без движения", "Реализовано", "Мониторинг использует фактическое время входа в стадию из stage_history; пороги SLA вынесены в настройки с историей и откатом.", "Проверить пороги на боевых воронках."),
        ("Без повторного касания", "Реализовано", "Правило использует загруженные звонки и встречи, фактическую последнюю активность и настраиваемый порог.", "Подтвердить бизнес-исключения по договорённостям с клиентом."),
        ("Отказы Спам/Дубль", "Частично", "Причина отказа может храниться в custom-полях; оценочные нарушения вынесены в настройки.", "Уточнить поле Bitrix24 и механизм ручной проверки."),
        ("SLA 1–6", "Реализовано", "Профиль SLA вынесен в Настройки → SLA; история стадий и активности дают фактические даты первого касания, движения и повторного контакта.", "Пункты, требующие экспертной оценки, принять с руководителями."),
        ("Задача из алерта", "Реализовано", "В мониторинге есть действие создания задачи через Bitrix24.", "Боевая приёмка прав webhook на crm.task.add."),
        ("Еженедельный AI-разбор", "Требует настройки", "Планировщик на понедельник 03:30 и хранилище AI-инсайтов реализованы.", "Нужен ключ LLM и боевые данные; сейчас AI-инсайтов 0."),
        ("AI-комментарий по сделке", "Реализовано", "Комментарий отображается в каждой строке. Есть детерминированная версия по CRM-фактам, fingerprint обновлений и LLM-генерация пакетами.", "Для LLM-комментариев нужен ключ и модель; сейчас заполнены базовые комментарии для 779 сделок."),
    ])

    doc.add_heading("6. Маркетинг, расходы и ROMI", level=1)
    add_callout(doc, "Решение по вопросу о ROMI", "Блок «ROMI по каналам» нужен по п. 3.6 ТЗ и возвращён на дашборд. Он не показывает демо-цифры: до появления расходов и привязанных поступлений 1С выводит честное пустое состояние.", fill=GREEN)
    add_matrix(doc, [
        ("Источники лидов", "Реализовано", "Дашборд → «Источники лидов»; разбивка по полю источника Bitrix24, фильтр источника.", "Проверить таксономию и UTM на боевых данных."),
        ("Расходы Директа", "Требует настройки", "Посуточная загрузка Direct Reports API, приведение к базе без НДС, разрез по юрлицу/периоду/кампании.", "Заполнить логины аккаунтов, проверить права и запустить выгрузку."),
        ("Ручные статьи расходов", "Реализовано", "Новое требование: Настройки → Расходы. Дата, юрлицо, статья, сумма, канал, кампания, комментарий; журнал и удаление.", "Заполняется пользователем."),
        ("Блок расходов", "Реализовано", "Дашборд → «Расходы по статьям»: сумма автоматического Директа и всех ручных статей; фильтр периода/юрлица.", "Появится после первой записи или выгрузки Директа."),
        ("ROMI по каналам", "Реализовано", "Дашборд и раздел ROMI. Формула: (база дохода − рекламный расход) / расход × 100%. В ROMI входят Direct и только ручные строки с флагом.", "Нужны факты Direct и связанные поступления 1С; иначе блок пуст. Не дублировать Direct вручную."),
        ("Маржа в ROMI", "Частично", "Выручка берётся только из разрешённых фактических поступлений 1С, не из успешной стадии Bitrix24.", "Источника себестоимости нет. До его появления это ROMI по выручке, а не по бухгалтерской марже."),
    ])

    doc.add_heading("7. Доступ и эксплуатация", level=1)
    add_matrix(doc, [
        ("Веб-панель", "Реализовано", "HTTPS, авторизация, закрытый API, Docker Compose, healthcheck, миграции, worker, backup-скрипты.", "Настроить операционный контроль бэкапов и алертов."),
        ("Роли: собственник/руководитель/менеджер", "Требует настройки", "RBAC реализован: собственник видит всё, руководитель ограничен отделом, менеджер — своими сделками; финансовые блоки менеджеру скрыты. Есть управление пользователями.", "В production отдельных пользователей 0. Создать учётные записи после заполнения сотрудников и отделов."),
        ("Telegram-бот", "Отложено", "Не подключён по последнему указанию Заказчика.", "Вернуть в объём после отдельного решения."),
        ("Calltouch", "Отложено", "Скрыт из интерфейса по решению Заказчика.", "Не включать в текущую приёмку."),
    ])

    doc.add_heading("8. Обязательная настройка перед приёмкой", level=1)
    doc.add_heading("Сотрудники и роли", level=2)
    for text in (
        "В настройках каждого входящего webhook Bitrix24 добавить право на краткий справочник пользователей — scope user_brief. Сейчас оба портала отклоняют user.get, поэтому система видит ID ответственного, но не получает ФИО.",
        "После изменения прав открыть «Интеграции», выполнить проверку обоих Bitrix24 и запустить пересчёт. Плановая синхронизация также обновит ФИО во всех ранее загруженных сделках, а не только в новых.",
        "Если право user_brief выдать нельзя, открыть «Настройки → Структура и планы → Сотрудники» и для каждого человека заполнить ФИО, портал box/cloud и ID пользователя Bitrix24. Ручное соответствие применяется к старым сделкам сразу после сохранения.",
        "Создать отделы и привязать к ним сотрудников и юрлица. После этого заполнить месячные планы по компании, отделам и сотрудникам: выручка, оплаты, сделки, звонки и встречи.",
        "В управлении пользователями создать учётные записи собственника, руководителей и менеджеров. Для руководителя выбрать отдел, для менеджера — соответствующего сотрудника.",
    ):
        add_number(doc, text)

    doc.add_heading("Три аккаунта Яндекса", level=2)
    for text in (
        "Для ЮО открыть «Интеграции → Яндекс · ЮО». OAuth-токен уже сохранён; заполнить логин клиента Яндекс Директа и числовой ID счётчика Метрики.",
        "Аналогично заполнить логин Директа и ID счётчика в карточках «Яндекс · ЦСВ» и «Яндекс · УрПАСЭ». OAuth-токены во всех трёх карточках уже присутствуют.",
        "Убедиться, что OAuth-пользователь имеет доступ на чтение каждого кабинета Директа и соответствующего счётчика Метрики. Логин Директа указывается именно для того рекламного клиента, данные которого должны попасть в юрлицо.",
        "Нажать «Проверить подключение» для каждой карточки. Исправить права или идентификаторы, если проверка не подтверждает одновременно Директ и Метрику.",
        "Запустить полный пересчёт. Проверить появление расходов в «Расходах по статьям», визитов и кампаний, а после загрузки поступлений 1С — значений в «ROMI по каналам».",
    ):
        add_number(doc, text)

    doc.add_heading("Остальные бизнес-настройки", level=2)
    for text in (
        "Оставить в панели ровно 11 согласованных воронок и привязать каждую к юрлицу и профилю SLA. Сейчас включено 15.",
        "Подтвердить стадии ожидаемой оплаты и финальную трактовку стадии «Заключение Контракта».",
        "Проверить боевую выгрузку 1С и разобрать журнал исключённых и несопоставленных поступлений.",
    ):
        add_number(doc, text)

    doc.add_heading("9. Приёмка нового блока расходов", level=1)
    for text in (
        "В «Настройки → Расходы» создаётся запись с датой, юрлицом, статьёй и положительной суммой.",
        "Строка удаляется из журнала только после подтверждения.",
        "Общий расход виден в «Расходах по статьям» за выбранные период и юрлицо.",
        "В ROMI попадает только ручная строка с флагом «Учитывать в ROMI»; для неё канал обязателен.",
        "Загруженный расход Директа появляется автоматически как отдельная статья; он не вводится вручную.",
        "При отсутствии расходов или связанной выручки ROMI не выдумывает число, а показывает пустое состояние.",
    ):
        add_bullet(doc, text)

    doc.add_heading("10. Техническая карта", level=1)
    add_two_col(doc, [
        ("Дашборд и ROMI", "DashboardPage.tsx; GET /api/dashboard/*; GET /api/romi/by-channel"),
        ("Расходы", "BusinessSettingsPage.tsx; /api/admin/expenses; /api/dashboard/expenses-by-article"),
        ("План-факт и отделы", "GET /api/dashboard/plan-fact; /departments; backend/app/services/plan_fact.py"),
        ("Роли и пользователи", "GET/POST/PUT/DELETE /api/admin/users; AppUser; RequireRole"),
        ("CRM и комментарии", "stage_history; crm_activities; deal_comments.py; Alembic 0010–0011 и 0013"),
        ("Настройки и интеграции", "/api/admin/business-settings; /api/integrations/settings; /bitrix/funnels"),
        ("Документация", "docs/ADMIN.md, DB_SCHEMA.md, ROMI_METHODOLOGY.md, ACCEPTANCE_CHECKLIST.md"),
    ])
    add_callout(doc, "Граница вывода", "Отчёт отражает состояние GitHub main и production на 05.09.2026. Пункты со статусом «Требует настройки» не считаются принятыми до заполнения доступов, структуры и планов и до появления боевых фактов. Секретные значения в документ не включены.", fill=LIGHT, trailing=False)

    core = doc.core_properties
    core.title = "Lawyer — отчёт о соответствии ТЗ"
    core.subject = "Сверка ТЗ, production и исходного кода"
    core.author = "Lawyer"
    core.keywords = "Lawyer, ТЗ, аудит, Bitrix24, 1С, ROMI"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
